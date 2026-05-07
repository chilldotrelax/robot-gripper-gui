from __future__ import annotations

import shutil
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree


ROOT = Path(__file__).resolve().parent
SRC = ROOT / "original.docx"
OUT = ROOT / "Final Report Template - repaired.docx"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
NS = {"w": W_NS, "wp": WP_NS}


def set_font(run_or_style, name: str, size: Pt | None = None) -> None:
    font = run_or_style.font
    font.name = name
    if size is not None:
        font.size = size
    rpr = run_or_style._element.rPr
    if rpr is None:
        rpr = run_or_style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), name)


def ensure_style(doc: Document, name: str, base: str = "normal"):
    try:
        return doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles[base]
        return style


def shade_paragraph(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Page ")
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def paragraph_has_page_break(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:br[@w:type='page']"))


def paragraph_has_drawing(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing"))


def configure_styles(doc: Document) -> None:
    normal = doc.styles["normal"]
    set_font(normal, "Times New Roman", Pt(11))
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.space_before = Pt(0)

    title = doc.styles["Title"]
    set_font(title, "Times New Roman", Pt(20))
    title.font.bold = True
    title.font.color.rgb = RGBColor(31, 55, 99)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(84)
    title.paragraph_format.space_after = Pt(16)

    subtitle = doc.styles["Subtitle"]
    set_font(subtitle, "Times New Roman", Pt(13))
    subtitle.font.italic = False
    subtitle.font.color.rgb = RGBColor(63, 63, 63)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)

    h1 = doc.styles["Heading 1"]
    set_font(h1, "Times New Roman", Pt(17))
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(31, 78, 121)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    set_font(h2, "Times New Roman", Pt(14))
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(36, 94, 92)
    h2.paragraph_format.space_before = Pt(15)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_font(h3, "Times New Roman", Pt(12))
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(63, 63, 63)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    toc_heading = ensure_style(doc, "TOC Heading")
    set_font(toc_heading, "Times New Roman", Pt(17))
    toc_heading.font.bold = True
    toc_heading.font.color.rgb = RGBColor(31, 78, 121)
    toc_heading.paragraph_format.space_before = Pt(18)
    toc_heading.paragraph_format.space_after = Pt(10)
    toc_heading.paragraph_format.keep_with_next = True

    caption = ensure_style(doc, "Caption")
    set_font(caption, "Times New Roman", Pt(9))
    caption.font.italic = True
    caption.font.color.rgb = RGBColor(89, 89, 89)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)

    body_label = ensure_style(doc, "Body Label")
    set_font(body_label, "Times New Roman", Pt(11))
    body_label.font.bold = True
    body_label.font.color.rgb = RGBColor(63, 63, 63)
    body_label.paragraph_format.space_before = Pt(6)
    body_label.paragraph_format.space_after = Pt(2)
    body_label.paragraph_format.keep_with_next = True

    spacer = ensure_style(doc, "Spacer")
    set_font(spacer, "Times New Roman", Pt(2))
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = 0.1

    code = ensure_style(doc, "Code Block")
    set_font(code, "Courier New", Pt(9))
    code.paragraph_format.left_indent = Inches(0.22)
    code.paragraph_format.right_indent = Inches(0.12)
    code.paragraph_format.space_before = Pt(1)
    code.paragraph_format.space_after = Pt(1)
    code.paragraph_format.line_spacing = 1.0


def normalize_paragraphs(doc: Document) -> None:
    long_heading_starts = (
        "Functioning robotic",
        "The gripper arm is",
        "Rather than servos",
        "This section focused on",
    )
    code_markers = (
        "#define",
        "CREATE variables:",
        "Setup:",
        "Main Loop:",
        "Open Gripper Function:",
        "IF ",
        "ELSE IF ",
        "Repeat ",
        "Delay ",
        "Turn ",
        "Set ",
        "Exit Loop",
        "Returns to",
        "user_input",
        "customlooprate",
        "……",
    )

    for i, p in enumerate(doc.paragraphs):
        text = p.text or ""
        stripped = text.strip()
        style_name = p.style.name if p.style is not None else ""

        # Make the existing TOC heading look like a heading without allowing the TOC field to include itself.
        if stripped == "Table of Contents":
            p.style = doc.styles["TOC Heading"]

        # Empty heading paragraphs were being pulled into the TOC. Keep their page breaks/spacer role only.
        elif not stripped:
            if style_name.startswith("Heading"):
                p.style = doc.styles["normal"]
            if not paragraph_has_drawing(p) and not paragraph_has_page_break(p):
                p.style = doc.styles["Spacer"]

        # The report used Heading 4 as the real third level, which made the TOC hierarchy brittle.
        elif style_name == "Heading 4":
            if stripped.startswith(long_heading_starts) or len(stripped) > 95:
                p.style = doc.styles["normal"]
            else:
                p.style = doc.styles["Heading 3"]

        # A few body paragraphs were accidentally promoted to Heading 2.
        elif style_name == "Heading 2" and (
            stripped.startswith(long_heading_starts)
            or (len(stripped) > 95 and "." in stripped)
        ):
            p.style = doc.styles["normal"]

        elif style_name == "Heading 2" and stripped in {
            "Deliverable - R1B",
            "Design & Justification:",
        }:
            p.style = doc.styles["Heading 3"]

        # Promote visually obvious subsection labels that were left as body text.
        if stripped in {"Summary", "Summary:", "Results:"}:
            p.style = doc.styles["Heading 3"]

        # Captions are real formatting, not body paragraphs.
        if stripped.startswith(("Figure:", "Figure:(", "Fig.", "Fig ")) or "Fig." in stripped:
            p.style = doc.styles["Caption"]

        # Short local labels should be bold labels, but not TOC headings.
        if (
            p.style.name == "normal"
            and 0 < len(stripped) <= 42
            and stripped.endswith(":")
            and not stripped.startswith(("http:", "https:"))
        ):
            p.style = doc.styles["Body Label"]

        if stripped.startswith(code_markers) or stripped.startswith("\t"):
            p.style = doc.styles["Code Block"]
            shade_paragraph(p, "F3F6F8")

        # Keep image-only paragraphs centered and separated from text.
        if paragraph_has_drawing(p) and not stripped:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(8)

        # Reset accidental paragraph drift so styles drive the document.
        if p.style.name not in {"Code Block"}:
            p.paragraph_format.left_indent = None
            p.paragraph_format.right_indent = None
            p.paragraph_format.first_line_indent = None

        for run in p.runs:
            # Preserve emphasis, underlines, and monospaced code feel; remove pasted font/size/color drift.
            if p.style.name == "Code Block":
                set_font(run, "Courier New", Pt(9))
            else:
                run.font.name = None
                run.font.size = None
                run.font.color.rgb = None


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.45)
        section.different_first_page_header_footer = True

        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        for run in list(p.runs):
            p._p.remove(run._r)
        add_page_field(p)
        p.style = doc.styles["normal"]
        for run in p.runs:
            set_font(run, "Times New Roman", Pt(9))
            run.font.color.rgb = RGBColor(89, 89, 89)


def patch_ooxml(docx_path: Path) -> None:
    tmp = Path(tempfile.mkdtemp(prefix="docx_repair_"))
    try:
        with zipfile.ZipFile(docx_path, "r") as z:
            z.extractall(tmp)

        document_xml = tmp / "word" / "document.xml"
        parser = etree.XMLParser(remove_blank_text=False)
        tree = etree.parse(str(document_xml), parser)
        root = tree.getroot()

        # Make the existing TOC refresh to Heading 1-3 only, avoiding stale Heading 4/body entries.
        for instr in root.xpath(".//w:instrText[contains(., 'TOC')]", namespaces=NS):
            instr.text = ' TOC \\o "1-3" \\h \\z \\u '

        # Set sane TOC tab stops within 6.5 in of usable text width.
        for sdt in root.xpath(
            ".//w:sdt[w:sdtPr/w:docPartObj/w:docPartGallery[@w:val='Table of Contents']]",
            namespaces=NS,
        ):
            for p in sdt.xpath(".//w:p", namespaces=NS):
                ppr = p.find("w:pPr", namespaces=NS)
                if ppr is None:
                    ppr = etree.Element(f"{{{W_NS}}}pPr")
                    p.insert(0, ppr)
                tabs = ppr.find("w:tabs", namespaces=NS)
                if tabs is None:
                    tabs = etree.SubElement(ppr, f"{{{W_NS}}}tabs")
                for tab in list(tabs):
                    tabs.remove(tab)
                tab = etree.SubElement(tabs, f"{{{W_NS}}}tab")
                tab.set(f"{{{W_NS}}}val", "right")
                tab.set(f"{{{W_NS}}}leader", "dot")
                tab.set(f"{{{W_NS}}}pos", "9360")

        # Floating images pasted from Docs used wrapNone/wrapSquare and allowed overlap. Keep the anchors,
        # but make them reserve vertical space so they do not sit on top of body text.
        for anchor in root.xpath(".//wp:anchor", namespaces=NS):
            anchor.set("allowOverlap", "0")
            for child in list(anchor):
                tag = etree.QName(child).localname
                if tag.startswith("wrap"):
                    anchor.remove(child)
            wrap = etree.Element(f"{{{WP_NS}}}wrapTopAndBottom")
            for idx, child in enumerate(anchor):
                if etree.QName(child).localname in {"docPr", "cNvGraphicFramePr", "graphic"}:
                    anchor.insert(idx, wrap)
                    break
            else:
                anchor.append(wrap)

        tree.write(str(document_xml), xml_declaration=True, encoding="UTF-8", standalone="yes")

        settings_xml = tmp / "word" / "settings.xml"
        if settings_xml.exists():
            stree = etree.parse(str(settings_xml), parser)
            sroot = stree.getroot()
        else:
            sroot = etree.Element(f"{{{W_NS}}}settings", nsmap={"w": W_NS})
            stree = etree.ElementTree(sroot)
        update_fields = sroot.find("w:updateFields", namespaces=NS)
        if update_fields is None:
            update_fields = etree.Element(f"{{{W_NS}}}updateFields")
            sroot.insert(0, update_fields)
        update_fields.set(f"{{{W_NS}}}val", "true")
        stree.write(str(settings_xml), xml_declaration=True, encoding="UTF-8", standalone="yes")

        with zipfile.ZipFile(docx_path, "w", compression=zipfile.ZIP_DEFLATED) as z:
            for p in tmp.rglob("*"):
                if p.is_dir():
                    continue
                z.write(p, p.relative_to(tmp).as_posix())
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def main() -> None:
    shutil.copyfile(SRC, OUT)
    doc = Document(str(OUT))
    configure_styles(doc)
    configure_sections(doc)
    normalize_paragraphs(doc)
    doc.save(str(OUT))
    patch_ooxml(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
