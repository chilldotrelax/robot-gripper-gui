from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


def main() -> None:
    path = Path(sys.argv[1])
    doc = Document(path)

    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sections={len(doc.sections)}")
    print("\nSECTIONS")
    for i, section in enumerate(doc.sections):
        print(
            i,
            "page",
            section.page_width,
            section.page_height,
            "margins",
            section.left_margin,
            section.right_margin,
            section.top_margin,
            section.bottom_margin,
            "header/footer",
            section.header_distance,
            section.footer_distance,
        )

    print("\nHEADERS")
    for i, section in enumerate(doc.sections):
        texts = []
        for p in section.header.paragraphs:
            txt = p.text.strip()
            if txt:
                texts.append(txt)
        print(i, " | ".join(texts) or "<empty>")

    print("\nFOOTERS")
    for i, section in enumerate(doc.sections):
        texts = []
        for p in section.footer.paragraphs:
            txt = p.text.strip()
            if txt:
                texts.append(txt)
        print(i, " | ".join(texts) or "<empty>")

    print("\nHEADINGS_AND_NEARBY")
    for idx, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style is not None else ""
        txt = " ".join(p.text.split())
        if style.startswith("Heading") or txt.lower().strip(": ") in {
            "summary",
            "requirements",
            "background",
            "budget",
            "appendix",
        }:
            print(f"\n#{idx} [{style}] {txt[:180]}")
            for j in range(max(0, idx - 2), min(len(doc.paragraphs), idx + 4)):
                q = doc.paragraphs[j]
                qstyle = q.style.name if q.style is not None else ""
                qtxt = " ".join(q.text.split())
                marker = "=>" if j == idx else "  "
                print(f"{marker} {j:03d} [{qstyle}] {qtxt[:180]}")

    print("\nTOC_LIKE")
    for idx, p in enumerate(doc.paragraphs):
        txt = " ".join(p.text.split())
        if "contents" in txt.lower() or "toc" in txt.lower():
            print(f"#{idx} [{p.style.name if p.style else ''}] {txt[:220]}")

    print("\nTABLES")
    for i, table in enumerate(doc.tables):
        dims = (len(table.rows), len(table.columns))
        sample = []
        for row in table.rows[:2]:
            sample.append(" || ".join(" ".join(cell.text.split())[:80] for cell in row.cells[:5]))
        print(f"table {i}: rows={dims[0]} cols={dims[1]} sample={sample}")


if __name__ == "__main__":
    main()
